import os
import re
import shutil
import sys
import subprocess
import yaml
from tqdm import tqdm
from docx import Document
import tkinter as tk
from tkinter import ttk
from tkinter import messagebox
from tqdm import tqdm

__version__ = "1.0.0"

def merge_pdfs(pdf1, pdf2, output_pdf):
    """
    Combine 2 PDFs.
    """
    doc1 = fitz.open(pdf1)
    doc2 = fitz.open(pdf2)
    merged_doc = fitz.open()
    merged_doc.insert_pdf(doc1)
    merged_doc.insert_pdf(doc2)
    merged_doc.save(output_pdf)
    print(f"PDF combined: {output_pdf}")

def latex_escape(s: str) -> str:
    """
    Handle special characters.
    """
    s = s.replace('\\', r'\textbackslash ')
    s = s.replace('&', r'\&')
    s = s.replace('%', r'\%')
    s = s.replace('$', r'\$')
    s = s.replace('#', r'\#')
    s = s.replace('_', r'\_')
    s = s.replace('~', r'\textasciitilde ')
    s = s.replace('^', r'\^{}')
    return s

def compile_local_texlive():
    """
    Compiling a PDF with Xelatex
    """
    os.chdir("./Latex")
    log_file = "./latex_compile.log"
    with open(log_file, "w") as log:
        subprocess.run(
            ["xelatex", "-interaction=nonstopmode", "main.tex"], 
            stdout=log, 
            stderr=log
        )
    os.chdir("..")

def Get_Basics(config):
    """
    Generating basic.tex from basic.yaml
    """
    latex_content = f"""
    \\newcommand{{\\course}}{{{config['Course_title']}}} 
    \\newcommand{{\\episode}}{{{config['Episode']}}} 
    \\newcommand{{\\week}}{{{config['week']}}}
    """
    with open("./Latex/Resources/page/basic.tex", "w", encoding="utf-8") as tex_file:
        tex_file.write(latex_content)

def convert_doc_to_latex(doc,RELATIONSHIP_NS,image_folder):
    """
    Generating output.tex from content.docx
    """
    latex_lines = []
    image_count = 0
    for table in doc.tables:
        for row in table.rows:
            if len(row.cells) < 2:
                continue
            result2 = ''
            for paragraph in row.cells[0].paragraphs:
                for run in paragraph.runs:
                    run_text = run.text
                    run_text = latex_escape(run_text) 
                    if run.bold:
                        result2 += r'\textbf{' + run_text + '}'
                    else:
                        result2 += run_text
            combined_text = re.sub(r"\s{2,}", " ", result2).strip()
            row_images = []
            for paragraph in row.cells[1].paragraphs:
                for run in paragraph.runs:
                    drawing_elems = run._element.xpath('.//*[local-name()="drawing"]')
                    if not drawing_elems:
                        continue
                    for drawing in drawing_elems:
                        blip_elems = drawing.xpath('.//*[local-name()="blip"]')
                        for blip in blip_elems:
                            rEmbed = blip.get(f'{{{RELATIONSHIP_NS}}}embed')
                            if not rEmbed:
                                continue
                            image_part = doc.part.related_parts[rEmbed]
                            image_count += 1
                            ext = image_part.content_type.split('/')[-1]
                            img_filename = f"{image_count}.{ext}"
                            img_path = os.path.join(image_folder, img_filename)
                            with open(img_path, "wb") as img_file:
                                img_file.write(image_part.blob)
                            row_images.append(img_filename)
            if len(row_images) == 0:
                latex_code = f"""
\\multicolumn{{3}}{{|c|}}{{%
    \\begin{{minipage}}{{\\linewidth}}
        \\centering
        \\vspace{{0.5em}}\\Large {combined_text}\\vspace{{0.5em}}
    \\end{{minipage}}
}}\\\\ \\hline          
"""
                latex_lines.append(latex_code)
            elif len(row_images) == 1:
                latex_code = f"""
\\centering \\steplist
&
{combined_text}
&
\\begin{{minipage}}[b]{{\\linewidth}}
    \\centering
    \\raisebox{{-.5\\height}}{{\\includegraphics[width=\\linewidth,height=0.3\\textheight,keepaspectratio]{{./Resources/figure_generated/{row_images[0]}}}}}
\\end{{minipage}}\\\\ \\hline
"""
                latex_lines.append(latex_code)
            elif len(row_images) == 2:
                latex_code1 = f"""
\\centering \\steplist
&
{combined_text}
&
\\begin{{minipage}}[b]{{\\linewidth}}
    \\centering
    \\raisebox{{-.5\\height}}{{\\includegraphics[width=\\linewidth,height=0.3\\textheight,keepaspectratio]{{./Resources/figure_generated/{row_images[0]}}}}}
\\end{{minipage}}\\\\ \\hline
"""
                latex_lines.append(latex_code1)
                latex_code2 = f"""
\\centering 
&
(Continued Image)
&
\\begin{{minipage}}[b]{{\\linewidth}}
    \\centering
    \\raisebox{{-.5\\height}}{{\\includegraphics[width=\\linewidth,height=0.3\\textheight,keepaspectratio]{{./Resources/figure_generated/{row_images[1]}}}}}
\\end{{minipage}}\\\\ \\hline
"""
                latex_lines.append(latex_code2)
            else:
                latex_code = f"""
\\centering \\steplist
&
{combined_text}
&
\\begin{{minipage}}[b]{{\\linewidth}}
    \\centering
    \\raisebox{{-.5\\height}}{{\\includegraphics[width=\\linewidth,height=0.3\\textheight,keepaspectratio]{{./Resources/figure_generated/{row_images[0]}}}}}
\\end{{minipage}}\\\\ \\hline
"""
                latex_lines.append(latex_code)
                for img in row_images[1:]:
                    latex_code = f"""
\\centering 
&
(Continued Image)
&
\\begin{{minipage}}[b]{{\\linewidth}}
    \\centering
    \\raisebox{{-.5\\height}}{{\\includegraphics[width=\\linewidth,height=0.3\\textheight,keepaspectratio]{{./Resources/figure_generated/{img}}}}}
\\end{{minipage}}\\\\ \\hline
"""
                    latex_lines.append(latex_code)
    return "\n".join(latex_lines)

def preprocess_doc(doc):
    """
    Combine Bold sections and clean the document.
    """
    for para in doc.paragraphs:
        merged_runs = []  
        for run in para.runs:
            text = run.text
            if text:
                current_bold = run.bold
                if merged_runs and merged_runs[-1][1] == current_bold:
                    merged_runs[-1] = (merged_runs[-1][0] + text, current_bold)
                else:
                    merged_runs.append((text, current_bold))
        p_element = para._element
        for child in list(p_element):
            p_element.remove(child)
        for text, bold in merged_runs:
            new_run = para.add_run(text)
            new_run.bold = bold
    return doc

def get_exe_dir():
    """
    Get the .exe path 
    """
    if getattr(sys, 'frozen', False):
        return os.path.dirname(sys.executable)
    else:
        return os.path.dirname(os.path.abspath(__file__))

def build():
    """
    Compiling the PDF and copy to higher level folder.
    """
    image_folder = "./Latex/Resources/figure_generated"
    os.makedirs(image_folder, exist_ok=True)
    RELATIONSHIP_NS = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'
    table_header = r"""
\begin{longtable}{|p{0.6cm}|m{5.5cm}|p{9cm}|}
\hline
Step & Description & Image \\ \hline
\endhead
"""
    table_footer = r"""
\end{longtable}
"""
    with open("basic.yaml", "r", encoding="utf-8") as file:
        config = yaml.safe_load(file)
    Get_Basics(config)
    tqdm.write(f"Compiling: {config['Course_title']} {config['Episode']} for {config['week']}")
    doc_path = "content.docx"
    doc = Document(doc_path)
    latex_lines = convert_doc_to_latex(preprocess_doc(doc),RELATIONSHIP_NS,image_folder)
    latex_output_path = "./Latex/Resources/page/output.tex"
    with open(latex_output_path, "w", encoding="utf-8") as f:
        f.write(table_header)
        f.writelines(latex_lines)
        f.write(table_footer)
    
    compile_local_texlive()
    pdf_source = os.path.join("Latex", "main.pdf")  
    pdf_dest = f"{config['week']} {config['Filename']}.pdf"  
    shutil.copyfile(pdf_source, pdf_dest)

def compile_in_subfolder(subfolder):
    """
    Runing build() in all subfolders.
    """
    original_dir = os.getcwd()
    try:
        os.chdir(subfolder)
        build()
    except Exception as e:
        print(f"{subfolder} error!: {e}")
    finally:
        os.chdir(original_dir)

def main():
    # Loading config 
    script_dir = get_exe_dir()
    subfolders_with_yaml = []
    for entry in os.scandir(script_dir):
        if entry.is_dir():
            basic_path = os.path.join(entry.path, "basic.yaml")
            if os.path.isfile(basic_path):
                try:
                    with open(basic_path, "r", encoding="utf-8") as f:
                        config = yaml.safe_load(f)
                    Course = config.get("Course_title", "Course_title")
                    Episode = config.get("Episode", "Episode")
                    week = config.get("week", "Week X")
                    subfolders_with_yaml.append({
                        "folder": entry.path,
                        "Course": Course,
                        "Episode": Episode,
                        "week": week
                    })
                except:
                    pass  
    if not subfolders_with_yaml:
        print("No basic.yaml Found")
        sys.exit(0)

    # GUI 
    root = tk.Tk()
    root.title("Chooes the project required")
    checkbox_vars = {}
    main_frame = ttk.Frame(root)
    main_frame.pack(padx=10, pady=10, fill=tk.BOTH, expand=True)
    row = 0
    for item in subfolders_with_yaml:
        folder_name = os.path.basename(item["folder"])
        display_text = f" {item['week']} {item['Episode']} in {folder_name} "
        var = tk.BooleanVar(value=True)  
        checkbox_vars[item["folder"]] = var
        chk = ttk.Checkbutton(main_frame, text=display_text, variable=var)
        chk.grid(row=row, column=0, sticky=tk.W, padx=5, pady=2)
        row += 1
    btn_frame = ttk.Frame(main_frame)
    btn_frame.grid(row=row, column=0, pady=10, sticky=tk.E)
    def select_all():
        for var in checkbox_vars.values():
            var.set(True)
    def select_none():
        for var in checkbox_vars.values():
            var.set(False)
    btn_select_all = ttk.Button(btn_frame, text="Select All", command=select_all)
    btn_select_all.pack(side=tk.LEFT, padx=5)
    btn_select_none = ttk.Button(btn_frame, text="Select None", command=select_none)
    btn_select_none.pack(side=tk.LEFT, padx=5)
    def on_confirm():
        selected_folders = [f for f in checkbox_vars if checkbox_vars[f].get()]
        if not selected_folders:
            messagebox.showinfo("Info", "No project selected!")
            return
        root.destroy()
        print("----------------------------------------------")
        print("Handout-Lx")
        print(f"Ver: {__version__}")
        print("----------------------------------------------")
        for folder in tqdm(selected_folders, desc="Processing folders"):
            compile_in_subfolder(folder)
        print(" ")
        print("All Done!")
    def on_cancel():
        root.destroy()
        print("User required exit.")
        sys.exit(0)
    ttk.Button(btn_frame, text="Confirm", command=on_confirm).pack(side=tk.LEFT, padx=5)
    ttk.Button(btn_frame, text="Cancel", command=on_cancel).pack(side=tk.LEFT)

    root.mainloop()

if __name__ == "__main__":
    main()
